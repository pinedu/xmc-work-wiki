"""
落档：把 docx 里的 N 部法规原文逐一写入 ~/wiki/concepts/某目录/

USAGE: 替换 SRC / DST_DIR / LAWS 三个变量即可。

来源：rhzy 体系 v2.1（2026-07-13）老板 3 次丢 docx 实战沉淀。
参考：本脚本来自 6 部法规批量落档，改改变量即可复用。
"""

import docx
from pathlib import Path

# ========== 必改 3 个变量 ==========
SRC = Path(r"C:\Users\rhzy\AppData\Local\hermes\cache\documents\doc_xxx.docx")
DST_DIR = Path(r"C:\Users\rhzy\wiki\concepts\某目录_全文版")
LAWS = [
    # 每部法规一行
    # {
    #     "id": 1,
    #     "name": "法规名",
    #     "law_id": "rhzy-bucket-law-xxx",
    #     "doc_num": "国务院令第 X 号",
    #     "publisher": "发文机关",
    #     "issued": "YYYY-MM-DD",
    #     "effective": "YYYY-MM-DD",
    #     "source_ref": "原文来源",
    #     "start": 3,         # 在 docx.paragraphs 中的起（包含）
    #     "end": 75,          # 在 docx.paragraphs 中的止（包含）
    #     "file_out": "法规名_全文.md",
    #     "type": "地方性法规 / 行政法规 / 部门规章",
    #     "summary_file": "../某目录/法规名.md",  # 已存在的摘要版路径
    # },
]

# ========== 模板代码（一般不用改） ==========
DST_DIR.mkdir(parents=True, exist_ok=True)

doc = docx.Document(str(SRC))
paragraphs = [p.text.strip() for p in doc.paragraphs]
total = len(paragraphs)

print(f"源文件: {SRC.name}")
print(f"段落总数: {total}")
print(f"目标目录: {DST_DIR}")
print(f"准备落档 {len(LAWS)} 部法规")
print("=" * 70)


def make_frontmatter(law: dict) -> str:
    """生成单部法规的 frontmatter + 元数据块"""
    return f"""---
title: {law['name']}
law_id: {law['law_id']}
doc_num: {law['doc_num']}
type: regulation/fulltext
publisher: {law['publisher']}
issued: {law['issued']}
effective: {law['effective']}
source_ref: {law['source_ref']}
created: 2026-07-13
source: {law.get('source_doc', '建设工程...docx')}
tags: [法规, 全文版, rhzy, 合规]
---

# {law['name']}

> **法规层级**：{law['type']}
> **文号**：`{law['doc_num']}`
> **发文机关**：{law['publisher']}
> **发布日期**：{law['issued']}
> **施行日期**：{law['effective']}
> **原文来源**：{law['source_ref']}
> **本档用途**：完整原文之一——含每条法条逐字全文。
> **配套档案**：摘要版 → [`{law['summary_file']}`]({law['summary_file']})（按 rhzy 业务方向提炼的关键条款）
> **本机路径索引**：本档落档于 `{DST_DIR / law['file_out']}`

---

"""


# 逐部落档
for law in LAWS:
    if law["end"] >= total:
        print(f"⚠️ [{law['id']}] {law['name']}: end 越界（total={total}），截断到 {total-1}")
        law["end"] = total - 1

    # 提取从 start 到 end 的所有非空段
    law_lines = [
        text for i in range(law['start'], law['end'] + 1)
        if (text := paragraphs[i])
    ]
    law_text = "\n\n".join(law_lines)

    # 写文件
    out_path = DST_DIR / law['file_out']
    content = make_frontmatter(law) + law_text
    out_path.write_text(content, encoding="utf-8")

    # 法条计数（用于报告）
    article_count = sum(
        1 for line in law_lines
        if line.startswith(("第", "制度")) and (
            "条" in line[:8] or "制度" in line[:6]
        )
    )
    print(f"✅ [{law['id']}] {law['name']}")
    print(f"   → {out_path}")
    print(f"   字节: {len(content):,} · 法条: ~{article_count}")

# 写索引 _index.md
INDEX_PATH = DST_DIR / "_index.md"
index_lines = [
    "---",
    f"title: {DST_DIR.name} · 索引",
    "type: index",
    "created: 2026-07-13",
    f"source: {SRC.name}",
    "tags: [法规, 全文版, rhzy, 合规]",
    "---",
    "",
    f"# {DST_DIR.name} · 全文版索引（{len(LAWS)} 部）",
    "",
    "> 本档收录 rhzy 项目 **{DST_DIR.name}** 的**逐条原文**完整版。",
    "> 已落档 {len(LAWS)} 部 = **摘要版** + **全文版** 双轨制（互相引用）。",
    "",
    "## 📋 法规清单",
    "",
    "| # | 法规名称 | 文号 | 类型 | 摘要版 | 全文版 |",
    "|---|---|---|---|---|---|",
]
for law in LAWS:
    index_lines.append(
        f"| {law['id']} | {law['name']} | {law['doc_num']} | {law['type']} | "
        f"[摘要]({law['summary_file']}) | [全文](./{law['file_out']}) |"
    )
index_lines.extend([
    "",
    f"**合计**：{len(LAWS)} 部 + 1 索引 · ~{sum(1 for _ in LAWS)} 条原文 · 详见各 .md",
    "",
    "## 🔗 联动业务主题",
    "",
    "本档覆盖以下主题（详见各法规 .md）：",
    "",
    "（按需补：农民工工资 / 招标投标 / 城乡规划 / 安全生产 / 建筑市场）",
    "",
    "## ⚠️ 引用规范",
    "",
    "- 引用条款时优先用**全文版**（找得到逐字原文）",
    "- 用"摘要版"作为速查入口",
    "- 任一引用均使用 4 层标识（业务 ID + 本机路径 + OSS URL + 锚点）",
])

INDEX_PATH.write_text("\n".join(index_lines), encoding="utf-8")
print(f"\n✅ 索引: {INDEX_PATH}")
print("\n" + "=" * 70)
print(f"落档完成：{len(LAWS)} 部法规 + 1 索引")
